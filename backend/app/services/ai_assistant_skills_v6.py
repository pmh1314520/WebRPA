"""WebRPA 小助手 - v6 联网搜索 & 网页研究 Skills

设计目标：补齐小助手/Agent 最关键的一块短板——**实时联网能力**。
此前小助手只能 http_get/fetch_page_html 抓"已知 URL"，无法"搜索未知信息"。
本模块让 AI 能像人一样：搜关键词 → 拿到结果列表 → 打开网页读正文 → 综合回答。

全部基于已内置依赖（httpx / beautifulsoup4 / html2text），无需任何 API Key：
  - web_search        DuckDuckGo 免密搜索（多端点容错），返回 标题/链接/摘要 列表
  - read_webpage      抓取网页并转成干净可读的正文（markdown），比原始 HTML 更适合 LLM 理解
  - research          一站式研究：搜索 → 自动打开前 N 条 → 汇总正文，供 AI 综合归纳

这些是 AI 自己直接调用的工具（后端真生效，不依赖前端编辑器在线）。
"""

from __future__ import annotations

import asyncio
import re
from html import unescape
from typing import Any
from urllib.parse import quote_plus, unquote, urlparse, parse_qs

from app.services.ai_assistant_skills import Skill, registry


_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
)


def _is_ad_url(href: str) -> bool:
    """识别 DuckDuckGo 的广告/赞助链接（y.js?ad_domain=... / ad_provider=...）。"""
    if not href:
        return True
    h = href.lower()
    return ("y.js" in h and ("ad_domain=" in h or "ad_provider=" in h)) or "/y.js?" in h


def _clean_ddg_url(href: str) -> str:
    """DuckDuckGo 的结果链接是重定向形式 //duckduckgo.com/l/?uddg=<编码真实URL>。
    解出真实 URL。若不是该形式则原样返回。广告链接返回空串以便过滤。"""
    if not href:
        return ""
    try:
        if href.startswith("//"):
            href = "https:" + href
        if _is_ad_url(href):
            return ""
        parsed = urlparse(href)
        if "duckduckgo.com" in (parsed.netloc or "") and parsed.path.startswith("/l/"):
            qs = parse_qs(parsed.query)
            real = (qs.get("uddg") or [""])[0]
            if real:
                real = unquote(real)
                return "" if _is_ad_url(real) else real
        return href
    except Exception:
        return href


def _strip_tags(html: str) -> str:
    """极简去标签（兜底用，正常走 BeautifulSoup）。"""
    text = re.sub(r"<[^>]+>", " ", html or "")
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


async def _ddg_html_search(query: str, max_results: int, timeout: int) -> list[dict[str, str]]:
    """用 DuckDuckGo HTML 端点搜索（免 API Key）。多端点容错。"""
    import httpx  # 已在 requirements 中

    endpoints = [
        ("https://html.duckduckgo.com/html/", "post"),
        ("https://lite.duckduckgo.com/lite/", "post"),
        ("https://duckduckgo.com/html/", "get"),
    ]
    headers = {
        "User-Agent": _UA,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        "Referer": "https://duckduckgo.com/",
    }
    last_err = ""
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True, headers=headers) as client:
        for url, method in endpoints:
            try:
                if method == "post":
                    resp = await client.post(url, data={"q": query, "kl": "wt-wt"})
                else:
                    resp = await client.get(url, params={"q": query, "kl": "wt-wt"})
                if resp.status_code != 200 or not resp.text:
                    last_err = f"HTTP {resp.status_code} @ {url}"
                    continue
                results = _parse_ddg_results(resp.text, max_results)
                if results:
                    return results
                last_err = f"无结果 @ {url}"
            except Exception as e:
                last_err = f"{type(e).__name__}: {e} @ {url}"
                continue
    if last_err:
        raise RuntimeError(last_err)
    return []


def _parse_ddg_results(html: str, max_results: int) -> list[dict[str, str]]:
    """解析 DuckDuckGo HTML / Lite 两种页面结构的搜索结果。"""
    results: list[dict[str, str]] = []
    try:
        from bs4 import BeautifulSoup  # 已在 requirements 中
    except Exception:
        return _parse_ddg_results_regex(html, max_results)

    soup = BeautifulSoup(html, "html.parser")

    # 结构 1：html.duckduckgo.com —— .result 容器
    for res in soup.select("div.result, div.web-result"):
        a = res.select_one("a.result__a")
        if not a:
            continue
        title = a.get_text(" ", strip=True)
        link = _clean_ddg_url(a.get("href", ""))
        snip_el = res.select_one(".result__snippet")
        snippet = snip_el.get_text(" ", strip=True) if snip_el else ""
        if title and link:
            results.append({"title": title, "url": link, "snippet": snippet})
        if len(results) >= max_results:
            return results

    if results:
        return results

    # 结构 2：lite.duckduckgo.com —— 表格里的 a.result-link
    for a in soup.select("a.result-link"):
        title = a.get_text(" ", strip=True)
        link = _clean_ddg_url(a.get("href", ""))
        if title and link:
            results.append({"title": title, "url": link, "snippet": ""})
        if len(results) >= max_results:
            break

    # 给 lite 结构补摘要（td.result-snippet）
    if results:
        snips = [td.get_text(" ", strip=True) for td in soup.select("td.result-snippet")]
        for i, r in enumerate(results):
            if i < len(snips):
                r["snippet"] = snips[i]
        return results

    return _parse_ddg_results_regex(html, max_results)


def _parse_ddg_results_regex(html: str, max_results: int) -> list[dict[str, str]]:
    """无 bs4 时的兜底正则解析。"""
    results: list[dict[str, str]] = []
    pattern = re.compile(
        r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>(.*?)</a>',
        re.IGNORECASE | re.DOTALL,
    )
    for m in pattern.finditer(html or ""):
        link = _clean_ddg_url(m.group(1))
        title = _strip_tags(m.group(2))
        if title and link:
            results.append({"title": title, "url": link, "snippet": ""})
        if len(results) >= max_results:
            break
    return results


def _clean_bing_url(href: str) -> str:
    """Bing 结果链接常是跳转形式 https://www.bing.com/ck/a?...&u=a1<base64url>。
    解出真实 URL；非跳转链接原样返回。"""
    if not href:
        return ""
    try:
        if "bing.com/ck/a" not in href:
            return href
        import base64
        from urllib.parse import urlparse, parse_qs
        u = (parse_qs(urlparse(href).query).get("u") or [""])[0]
        if u.startswith("a1"):
            b = u[2:]
            b += "=" * (-len(b) % 4)
            decoded = base64.urlsafe_b64decode(b).decode("utf-8", "ignore")
            if decoded.startswith("http"):
                return decoded
        return href
    except Exception:
        return href


def _parse_bing_results(html: str, max_results: int) -> list[dict[str, str]]:
    """解析 Bing 搜索结果页（li.b_algo）。"""
    results: list[dict[str, str]] = []
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return results
    soup = BeautifulSoup(html, "html.parser")
    for li in soup.select("li.b_algo"):
        a = li.select_one("h2 a") or li.select_one("a")
        if not a:
            continue
        link = _clean_bing_url(a.get("href", ""))
        title = a.get_text(" ", strip=True)
        if not (link and title) or not link.startswith("http") or "bing.com/ck/a" in link:
            continue
        snip_el = li.select_one(".b_caption p") or li.select_one("p")
        snippet = snip_el.get_text(" ", strip=True) if snip_el else ""
        results.append({"title": title, "url": link, "snippet": snippet})
        if len(results) >= max_results:
            break
    return results


async def _bing_html_search(query: str, max_results: int, timeout: int) -> list[dict[str, str]]:
    """用 Bing 网页端搜索（免 API Key），作为 DuckDuckGo 的兜底引擎。"""
    import httpx
    headers = {
        "User-Agent": _UA,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9,zh-CN;q=0.8",
        # 关键：带上市场/语言 cookie，否则 Bing 可能返回不含 b_algo 的市场首页
        "Cookie": "SRCHHPGUSR=SRCHLANG=en; _EDGE_S=mkt=en-us",
    }
    url = "https://www.bing.com/search"
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True, headers=headers) as client:
        resp = await client.get(url, params={"q": query})
        if resp.status_code != 200 or not resp.text:
            raise RuntimeError(f"Bing HTTP {resp.status_code}")
        return _parse_bing_results(resp.text, max_results)


async def _search(query: str, max_results: int, timeout: int) -> tuple[list[dict[str, str]], str]:
    """统一搜索入口：DuckDuckGo 优先，失败/空结果时自动切 Bing。
    返回 (结果列表, 实际使用的引擎名)。两个引擎都失败则抛出最后的异常。"""
    errors: list[str] = []
    # 1) DuckDuckGo
    try:
        ddg = await _ddg_html_search(query, max_results, timeout)
        if ddg:
            return ddg, "duckduckgo"
    except Exception as e:
        errors.append(f"ddg: {e}")
    # 2) Bing 兜底
    try:
        bing = await _bing_html_search(query, max_results, timeout)
        if bing:
            return bing, "bing"
    except Exception as e:
        errors.append(f"bing: {e}")
    if errors:
        raise RuntimeError("；".join(errors))
    return [], "none"


# =============================================================================
# Skill 处理函数
# =============================================================================

async def skill_web_search(
    query: str,
    max_results: int = 8,
    timeout: int = 15,
    **_: Any,
) -> dict[str, Any]:
    """联网搜索（DuckDuckGo，免 API Key）。返回标题/链接/摘要列表。

    用于查实时信息：最新版本、官方文档、报错解法、新闻、教程等。
    拿到链接后可用 read_webpage 打开读正文，或直接基于摘要回答。
    """
    q = (query or "").strip()
    if not q:
        return {"error": "query 不能为空"}
    try:
        n = max(1, min(int(max_results or 8), 20))
    except Exception:
        n = 8
    try:
        results, engine = await _search(q, n, int(timeout or 15))
    except Exception as e:
        return {"error": f"搜索失败：{e}", "query": q, "results": []}
    if not results:
        return {
            "query": q,
            "count": 0,
            "results": [],
            "note": "未搜到结果，可换个关键词或更具体的措辞重试。",
        }
    return {"query": q, "count": len(results), "results": results, "engine": engine}


async def skill_read_webpage(
    url: str,
    max_chars: int = 20000,
    timeout: int = 20,
    include_links: bool = False,
    **_: Any,
) -> dict[str, Any]:
    """抓取网页并转成干净可读的正文（markdown）。

    比 fetch_page_html 的原始 HTML 更适合 LLM 理解：自动去掉 script/style/nav/footer，
    保留标题/段落/列表/表格结构。用于"读文档、读文章、读报错页面"。
    """
    u = (url or "").strip()
    if not u:
        return {"error": "url 不能为空"}
    if not re.match(r"^https?://", u, re.IGNORECASE):
        u = "https://" + u
    try:
        import httpx
    except ImportError:
        return {"error": "需要 httpx：pip install httpx"}

    headers = {
        "User-Agent": _UA,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }
    try:
        async with httpx.AsyncClient(timeout=int(timeout or 20), follow_redirects=True, headers=headers) as client:
            resp = await client.get(u)
            final_url = str(resp.url)
            ctype = resp.headers.get("content-type", "")
            raw = resp.text
    except Exception as e:
        return {"error": f"抓取失败：{e}", "url": u}

    if resp.status_code >= 400:
        return {"error": f"HTTP {resp.status_code}", "url": final_url, "status_code": resp.status_code}

    # 非 HTML（如 JSON / 纯文本）直接返回文本
    if "html" not in ctype.lower() and not raw.lstrip().startswith("<"):
        text = raw
        truncated = len(text) > max_chars
        return {
            "url": final_url,
            "title": "",
            "content": text[:max_chars],
            "truncated": truncated,
            "content_type": ctype,
        }

    title = ""
    text = ""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        if soup.title and soup.title.string:
            title = soup.title.string.strip()
        # 去掉噪音标签
        for tag in soup(["script", "style", "noscript", "iframe", "svg", "nav", "footer", "header", "form"]):
            tag.decompose()
        # 优先正文容器
        main = (
            soup.find("article")
            or soup.find("main")
            or soup.find(attrs={"role": "main"})
            or soup.body
            or soup
        )
        cleaned_html = str(main)
        try:
            import html2text
            h = html2text.HTML2Text()
            h.ignore_links = not include_links
            h.ignore_images = True
            h.body_width = 0
            text = h.handle(cleaned_html)
        except Exception:
            text = main.get_text("\n", strip=True)
    except Exception:
        text = _strip_tags(raw)

    # 压缩多余空行
    text = re.sub(r"\n{3,}", "\n\n", text or "").strip()
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars] + f"\n\n…（正文过长，已截断，仅保留前 {max_chars} 字）"

    return {
        "url": final_url,
        "title": title,
        "content": text,
        "truncated": truncated,
        "content_type": ctype,
    }


async def skill_research(
    query: str,
    max_pages: int = 3,
    max_chars_per_page: int = 6000,
    timeout: int = 20,
    **_: Any,
) -> dict[str, Any]:
    """一站式联网研究：搜索 → 并发打开前 N 条结果 → 汇总正文。

    适合"调研型"提问（如"对比 X 和 Y""某框架最新用法""某报错怎么解决"）。
    返回每条来源的 标题/链接/正文摘要，AI 据此综合归纳并给出带来源的答案。
    """
    q = (query or "").strip()
    if not q:
        return {"error": "query 不能为空"}
    try:
        k = max(1, min(int(max_pages or 3), 5))
    except Exception:
        k = 3

    try:
        results, engine = await _search(q, max(k, 6), int(timeout or 20))
    except Exception as e:
        return {"error": f"搜索失败：{e}", "query": q, "sources": []}
    if not results:
        return {"query": q, "sources": [], "note": "未搜到结果，可换关键词重试。"}

    targets = results[:k]

    async def _read_one(item: dict[str, str]) -> dict[str, Any]:
        page = await skill_read_webpage(
            url=item["url"],
            max_chars=int(max_chars_per_page or 6000),
            timeout=int(timeout or 20),
        )
        return {
            "title": item.get("title") or page.get("title") or "",
            "url": item["url"],
            "snippet": item.get("snippet", ""),
            "content": page.get("content", "") if not page.get("error") else "",
            "error": page.get("error", ""),
        }

    sources = await asyncio.gather(*[_read_one(it) for it in targets], return_exceptions=True)
    clean_sources: list[dict[str, Any]] = []
    for s in sources:
        if isinstance(s, dict):
            clean_sources.append(s)

    ok = sum(1 for s in clean_sources if s.get("content"))
    return {
        "query": q,
        "engine": engine,
        "source_count": len(clean_sources),
        "readable_count": ok,
        "sources": clean_sources,
        "all_results": results,
        "note": "请综合各来源正文归纳回答，并在结尾用 [标题](URL) 形式标注引用来源。",
    }


# =============================================================================
# 注册
# =============================================================================

def _register_v6() -> None:
    registry.register(Skill(
        name="web_search",
        description=(
            "联网搜索（DuckDuckGo，免 API Key，后端直连）。返回 标题/链接/摘要 列表。"
            "用于查实时/未知信息：库的最新版本、官方文档、报错解法、API 用法、新闻、教程等。"
            "拿到链接后可用 read_webpage 读正文，或直接基于摘要回答。"
            "**当用户问到你不确定或可能已过时的信息时，主动搜一下再答，别凭记忆瞎猜。**"
        ),
        parameters={
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "搜索关键词，越具体越好"},
                "max_results": {"type": "integer", "default": 8, "description": "返回结果数（1-20）"},
                "timeout": {"type": "integer", "default": 15},
            },
            "required": ["query"],
        },
        handler=skill_web_search,
    ))

    registry.register(Skill(
        name="read_webpage",
        description=(
            "抓取网页并转成干净可读的正文（markdown，自动去 script/style/nav/广告）。"
            "比 fetch_page_html 的原始 HTML 更适合理解长文。用于读官方文档、文章、报错页等。"
            "通常配合 web_search：先搜到链接，再用本工具读正文。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "网页地址"},
                "max_chars": {"type": "integer", "default": 20000, "description": "正文最多保留字数"},
                "timeout": {"type": "integer", "default": 20},
                "include_links": {"type": "boolean", "default": False, "description": "是否在正文中保留超链接"},
            },
            "required": ["url"],
        },
        handler=skill_read_webpage,
    ))

    registry.register(Skill(
        name="research",
        description=(
            "一站式联网研究：搜索 → 并发打开前 N 条结果 → 汇总各页正文。"
            "适合调研型问题（对比/选型/最新用法/疑难报错）。返回多来源正文，"
            "你需综合归纳并在结尾用 [标题](URL) 标注引用来源。比手动 web_search + 多次 read_webpage 更省事。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "研究主题/问题"},
                "max_pages": {"type": "integer", "default": 3, "description": "打开并读取的结果数（1-5）"},
                "max_chars_per_page": {"type": "integer", "default": 6000, "description": "每页正文保留字数"},
                "timeout": {"type": "integer", "default": 20},
            },
            "required": ["query"],
        },
        handler=skill_research,
    ))

    registry.register(Skill(
        name="download_file",
        description=(
            "下载任意网络文件到本地（流式，支持 PDF/数据集/安装包等）。"
            "save_path 留空则自动存到临时目录。常配合 read_document 使用：先下载再读正文。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "文件地址"},
                "save_path": {"type": "string", "description": "本地保存路径，留空自动生成"},
                "timeout": {"type": "integer", "default": 60},
                "max_mb": {"type": "integer", "default": 100, "description": "大小上限(MB)，超过则中止"},
            },
            "required": ["url"],
        },
        handler=skill_download_file,
    ))

    registry.register(Skill(
        name="read_document",
        description=(
            "读取文档为纯文本，支持 PDF / Word(docx) / Excel(xlsx/xls) / CSV / TXT 等。"
            "source 可以是本地路径或网络 URL（URL 会先自动下载再解析）。"
            "补齐 read_webpage 读不了二进制文档的短板——读下载来的 PDF 论文/手册、本地 Excel 报表等用它。"
        ),
        parameters={
            "type": "object",
            "properties": {
                "source": {"type": "string", "description": "本地路径或网络 URL"},
                "max_chars": {"type": "integer", "default": 20000, "description": "正文最多保留字数"},
                "timeout": {"type": "integer", "default": 60},
            },
            "required": ["source"],
        },
        handler=skill_read_document,
    ))


# =============================================================================
# 文件下载 & 文档读取（让研究能力延伸到 PDF/Word/Excel）
# =============================================================================

async def skill_download_file(
    url: str,
    save_path: str = "",
    timeout: int = 60,
    max_mb: int = 100,
    **_: Any,
) -> dict[str, Any]:
    """下载任意网络文件到本地（流式，适合 PDF/数据集/安装包等）。

    save_path 留空则自动按 URL 文件名存到系统临时目录。返回本地路径与大小。
    """
    u = (url or "").strip()
    if not u:
        return {"error": "url 不能为空"}
    if not re.match(r"^https?://", u, re.IGNORECASE):
        u = "https://" + u
    try:
        import httpx
    except ImportError:
        return {"error": "需要 httpx：pip install httpx"}

    import os
    import tempfile
    from urllib.parse import urlparse as _up

    dest = (save_path or "").strip()
    if not dest:
        name = os.path.basename(_up(u).path) or "download.bin"
        if "." not in name:
            name += ".bin"
        dest = os.path.join(tempfile.gettempdir(), "webrpa_dl", name)
    os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)

    limit = max(1, int(max_mb or 100)) * 1024 * 1024
    headers = {"User-Agent": _UA}
    written = 0
    try:
        async with httpx.AsyncClient(timeout=int(timeout or 60), follow_redirects=True, headers=headers) as client:
            async with client.stream("GET", u) as resp:
                if resp.status_code >= 400:
                    return {"error": f"HTTP {resp.status_code}", "url": u}
                with open(dest, "wb") as f:
                    async for chunk in resp.aiter_bytes(65536):
                        written += len(chunk)
                        if written > limit:
                            f.close()
                            try:
                                os.remove(dest)
                            except Exception:
                                pass
                            return {"error": f"文件超过 {max_mb}MB 上限，已中止下载"}
                        f.write(chunk)
                ctype = resp.headers.get("content-type", "")
    except Exception as e:
        return {"error": f"下载失败：{e}", "url": u}

    return {
        "success": True,
        "path": dest,
        "size_bytes": written,
        "size_human": f"{written / 1024 / 1024:.2f} MB" if written >= 1024 * 1024 else f"{written / 1024:.1f} KB",
        "content_type": ctype,
    }


def _extract_document_text(path: str, ext: str, max_chars: int) -> str:
    """从本地文档文件提取纯文本（pdf/docx/xlsx/xls/csv/txt 等）。"""
    ext = (ext or "").lower()

    def _clip(t: str) -> str:
        t = t or ""
        if len(t) > max_chars:
            return t[:max_chars] + f"\n…（内容过长，已截断，仅保留前 {max_chars} 字）"
        return t

    # 纯文本类
    if ext in (".txt", ".md", ".markdown", ".csv", ".tsv", ".log", ".json", ".xml",
               ".yaml", ".yml", ".ini", ".html", ".htm"):
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "utf-16", "latin-1"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return _clip(f.read())
            except Exception:
                continue
        with open(path, "rb") as f:
            return _clip(f.read().decode("utf-8", errors="ignore"))

    if ext == ".pdf":
        try:
            import pdfplumber
            parts: list[str] = []
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    if i >= 50:
                        parts.append("…（页数过多，仅解析前 50 页）")
                        break
                    parts.append(page.extract_text() or "")
            return _clip("\n".join(parts))
        except Exception:
            from pypdf import PdfReader
            reader = PdfReader(path)
            parts = [(p.extract_text() or "") for p in reader.pages[:50]]
            return _clip("\n".join(parts))

    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return _clip("\n".join(parts))

    if ext == ".xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        out: list[str] = []
        for ws in wb.worksheets:
            out.append(f"# 工作表: {ws.title}")
            for ri, row in enumerate(ws.iter_rows(values_only=True)):
                if ri >= 500:
                    out.append("…（行数过多，仅解析前 500 行）")
                    break
                out.append("\t".join("" if c is None else str(c) for c in row))
        wb.close()
        return _clip("\n".join(out))

    if ext == ".xls":
        import xlrd
        wb = xlrd.open_workbook(path)
        out = []
        for ws in wb.sheets():
            out.append(f"# 工作表: {ws.name}")
            for ri in range(min(ws.nrows, 500)):
                out.append("\t".join(str(v) for v in ws.row_values(ri)))
        return _clip("\n".join(out))

    if ext == ".doc":
        raise RuntimeError("暂不支持旧版 .doc（请另存为 .docx）")

    # 兜底当文本
    with open(path, "rb") as f:
        return _clip(f.read().decode("utf-8", errors="ignore"))


async def skill_read_document(
    source: str,
    max_chars: int = 20000,
    timeout: int = 60,
    **_: Any,
) -> dict[str, Any]:
    """读取文档为纯文本，支持 PDF / Word(docx) / Excel(xlsx/xls) / CSV / TXT 等。

    source 可以是**本地路径**，也可以是**网络 URL**（会先自动下载到临时目录再解析）。
    用于"读下载来的 PDF 论文/手册""读本地 Excel 报表"等，补齐 read_webpage 读不了二进制文档的短板。
    """
    import os
    import tempfile

    src = (source or "").strip()
    if not src:
        return {"error": "source 不能为空"}

    tmp_to_clean = ""
    try:
        if re.match(r"^https?://", src, re.IGNORECASE):
            dl = await skill_download_file(url=src, timeout=timeout)
            if dl.get("error"):
                return {"error": f"下载失败：{dl['error']}", "source": src}
            path = dl["path"]
            tmp_to_clean = path
        else:
            path = src
            if not os.path.exists(path):
                return {"error": f"文件不存在：{path}"}

        ext = os.path.splitext(path)[1].lower()
        try:
            text = _extract_document_text(path, ext, int(max_chars or 20000))
        except Exception as e:
            return {"error": f"解析失败：{type(e).__name__}: {e}", "source": src, "ext": ext}

        return {
            "source": src,
            "ext": ext,
            "chars": len(text),
            "content": text,
        }
    finally:
        # 只清理我们临时下载的文件，不动用户的本地文件
        if tmp_to_clean and tempfile.gettempdir() in tmp_to_clean:
            try:
                os.remove(tmp_to_clean)
            except Exception:
                pass


_register_v6()
